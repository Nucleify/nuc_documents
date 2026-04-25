from enum import Enum
from io import BytesIO
from typing import BinaryIO, Callable

import tabula
from borb import pdf
from borb.pdf.layout_element.table.table_util import TableUtil
from docx import Document
from fastapi import FastAPI, Response, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from odf.opendocument import load
from odf.table import Table, TableCell, TableRow
from odf.teletype import extractText
from pandas import DataFrame, concat, read_csv, read_json, read_xml

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class Formats(Enum):
    CSV = "csv"
    XML = "xml"
    JSON = "json"
    PDF = "pdf"


class ValidExtensions(Enum):
    CSV = "csv"
    XML = "xml"
    JSON = "json"
    PDF = "pdf"
    DOC = "doc"
    DOCX = "docx"
    ODT = "odt"


MAX_PDF_COLUMNS = 8
MAX_PDF_ROWS = 250
MAX_PDF_CELL_CHARS = 32


def _truncate_text(value: object, limit: int = MAX_PDF_CELL_CHARS) -> str:
    text = "" if value is None else str(value)
    return text if len(text) <= limit else f"{text[: limit - 1]}…"


def to_pdf(df: DataFrame) -> Response:
    rows_limit = min(MAX_PDF_ROWS, len(df.index))
    column_candidates = [MAX_PDF_COLUMNS, 6, 4, 3, 2, 1]
    chars_candidates = [MAX_PDF_CELL_CHARS, 24, 18, 14, 10]

    for column_limit in column_candidates:
        for chars_limit in chars_candidates:
            prepared_df = df.fillna("").copy()
            prepared_df = prepared_df.iloc[:rows_limit, :column_limit]
            prepared_df.columns = [_truncate_text(column, chars_limit) for column in prepared_df.columns]
            prepared_df = prepared_df.astype(str).apply(lambda col: col.map(lambda v: _truncate_text(v, chars_limit)))

            _pdf = pdf.Document()
            page = pdf.Page(width=842, height=595)
            _pdf.append_page(page)
            layout = pdf.SingleColumnLayout(page)

            tabular = [prepared_df.keys().tolist()]
            for row in prepared_df.values.tolist():
                tabular.append(row)

            try:
                table = TableUtil.from_2d_data(tabular)
                layout.append_layout_element(table)

                buffered = BytesIO()
                pdf.PDF.write(_pdf, buffered)
                return Response(
                    content=buffered.getvalue(),
                    headers={"content-disposition": "attachment; filename=result.pdf"},
                    media_type="application/pdf",
                )
            except AssertionError:
                continue

    fallback_pdf = pdf.Document()
    fallback_page = pdf.Page(width=842, height=595)
    fallback_pdf.append_page(fallback_page)
    fallback_layout = pdf.SingleColumnLayout(fallback_page)
    fallback_table = TableUtil.from_2d_data(
        [["PDF conversion fallback"], ["Dataset too wide; exported in compact preview mode."]]
    )
    fallback_layout.append_layout_element(fallback_table)

    buffered = BytesIO()
    pdf.PDF.write(fallback_pdf, buffered)
    return Response(
        content=buffered.getvalue(),
        headers={"content-disposition": "attachment; filename=result.pdf"},
        media_type="application/pdf",
    )


FUNCTIONS: dict[Formats, Callable[[DataFrame], Response]] = {
    Formats.CSV: lambda df: Response(content=df.to_csv(index=False), media_type="text/csv"),
    Formats.JSON: lambda df: Response(content=df.to_json(orient="records"), media_type="application/json"),
    Formats.XML: lambda df: Response(content=df.to_xml(index=False), media_type="application/xml"),
    Formats.PDF: to_pdf,
}


def from_odt(file: BinaryIO) -> DataFrame:
    """
    Extracts tables from an ODT file and returns a concatenated DataFrame.
    """
    doc = load(file)
    dfs = []
    for table in doc.getElementsByType(Table):
        data = []
        for row in table.getElementsByType(TableRow):
            row_data = []
            for cell in row.getElementsByType(TableCell):
                cell_text = extractText(cell)
                row_data.append(cell_text)
            data.append(row_data)
        if data:
            header = data[0]
            df = DataFrame(data[1:], columns=header)
            dfs.append(df)
    if not dfs:
        return DataFrame()
    return concat(dfs, ignore_index=True)


def from_pdf(file: BinaryIO) -> DataFrame:
    """
    Extracts tables from a PDF file and returns a concatenated DataFrame.
    """
    try:
        tables = tabula.read_pdf(file, pages="all", multiple_tables=True)
    except Exception:
        return DataFrame()
    if not tables:
        return DataFrame()
    return concat(tables, ignore_index=True)


def from_docx(file: BinaryIO) -> DataFrame:
    """
    Extracts tables from a DOCX file and returns a concatenated DataFrame.
    """
    doc = Document(file)
    dfs = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        if data:
            header = data[0]
            df = DataFrame(data[1:], columns=header)
            dfs.append(df)
    if not dfs:
        return DataFrame()
    return concat(dfs, ignore_index=True)


@app.post("/")
def transform(format: Formats, file: UploadFile):
    """
    Transforms given file to another format
    Parameters
    ----------
    format:
        Destination format to use
    file:
        File to transform. Requires valid extension
    """
    extension = file.filename.split(".")[-1] if file.filename else ""

    if extension == "csv":
        df = read_csv(file.file)
    elif extension == "xml":
        df = read_xml(file.file)
    elif extension == "json":
        df = read_json(file.file)
    elif extension == "docx":
        df = from_docx(file.file)
    elif extension == "odt":
        df = from_odt(file.file)
    elif extension == "pdf":
        df = from_pdf(file.file)
    elif extension == "doc":
        return Response(
            content="ERROR: .doc format is not supported for in-memory structured data extraction without external tools.",
            status_code=400,
        )
    else:
        return Response(content=f"Unsupported file extension: {extension}", status_code=400)

    if df.empty:
        return Response(content="Could not extract data from file.", status_code=400)

    if func := FUNCTIONS.get(format):
        try:
            return func(df)
        except AssertionError:
            return Response(
                content=(
                    "Could not render a PDF table for this dataset (too wide). "
                    "Try converting to CSV, JSON, or XML."
                ),
                status_code=400,
            )
    return Response(content=f"Unsupported output format: {format.value}", status_code=400)
